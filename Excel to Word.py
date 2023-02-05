

import pandas as pd
import numpy as np
from tabulate import tabulate
import docx
from docx import Document
from docx.oxml import OxmlElement, ns
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import glob

""" ########################################################################################################################## """
""" Import Survey Responses """
""" ########################################################################################################################## """

# Set folder path with all participants' answers (responses per participant are in a separate Excel file)
path_survey_response = glob.glob(r'T:\StaffFolders\JFang\Project\Responses\*.xlsx')

dict_responses = {}
participant_name = ['Participant 1', 'Participant 2', 'Participant 3']

# Stores each participants' responses are stored in a DataFrame in a Dictionary
for participant, path in zip(participant_name, path_survey_response):
    dict_responses[participant] = pd.read_excel(path)

for df in dict_responses.values():
    print(tabulate(df.head(2), headers='keys', tablefmt='psql', numalign='right', showindex=False))

# Stores each participants' responses in a single DataFrame
df_responses_concat = pd.concat(dict_responses.values(), ignore_index=True)


""" ########################################################################################################################## """
""" Import Original Word Document with Question Descriptions """
""" ########################################################################################################################## """

path_survey_questions = docx.Document(r'T:\StaffFolders\JFang\Project\Survey Instructions.docx')

text_questions = []
for paragraph in path_survey_questions.paragraphs:
    print(paragraph.text)
    text_questions.append(paragraph.text)

list_questions = pd.DataFrame(text_questions, columns='Question')

list_questions2 = list_questions[21:].copy()  # Removes all the introductory paragraphs in the word document
list_questions3 = list_questions2[list_questions2['Question'] != ''].copy()  # Removes blank rows

# Questions are in the format of 'Column A: Describe...', so need to separate out the Question Number/Column Letter with the
# actual question itself
list_questions3['Question Number'] = list_questions3['Question'].str.split(': ').str[0]  # Question Number is actually Column
# Letter (e.g. Column A instead of Question 1)
list_questions3['Question'] = list_questions3['Question'].str.split(': ').str[1]

df_questions = list_questions3[['Question Number', 'Question']].copy()

print(tabulate(df_questions.head(2), headers='keys', tablefmt='psql', numalign='right', showindex=False))


""" ########################################################################################################################## """
""" Preview/Print Responses """
""" ########################################################################################################################## """

# Makes sure responses are printed in the desired format before writing to a Word Document


# Assigns each Column Number a number for referencing by index later on
a = 0
b = 1
c = 2
d = 3
e = 4
f = 5
g = 6
h = 7
i = 8
j = 9
k = 10
l = 11
m = 12
n = 13
o = 14
p = 15
q = 16
r = 17
s = 18
t = 19
u = 20
v = 21
w = 22
x = 23
y = 24
z = 25
aa = 26
ab = 27
ac = 28
ad = 29
ae = 30
af = 31
ag = 32
ah = 33
ai = 34
aj = 35
ak = 36
al = 37
am = 38
an = 39
ao = 40
ap = 41
aq = 42
ar = 43
AS = 44
at = 45
au = 46
av = 47
aw = 48
ax = 49
ay = 50
az = 51
ba = 52
bb = 53
bc = 54
bd = 55
be = 56
bf = 57
bg = 58
bh = 59
bi = 60
bj = 61
bk = 62
bl = 63
bm = 64
bn = 65
bo = 66
bp = 67
bq = 68
br = 69
bs = 70
bt = 71
bu = 72
bv = 73
bw = 74
bx = 75
by = 76
bz = 77


""" Print All Responses per Unit """

# Word Document will not necessarily print all Questions/Columns in sequential order. Only certain questions will be
# printed/pasted into a given section of the final Word Document.

questions_in_base_section = [k, l, m, n, o]

for question_num in questions_in_base_section:

    # Question
    print('')
    print(df_questions.iloc[question_num, 0] + ' - ' + df_questions.iloc[question_num, 1])  # Prints Question Number and Question
    print('')

    for participant, participant_response in dict_responses.items():
        print(participant)  # Prints name of Participant

        for unit in np.arange(len(participant_response)):
            print(participant_response.iloc[unit, 3])  # Prints name of Unit, which is provided in the 4th column
            print(participant_response.iloc[unit, question_num])  # Response
        print('')


""" Print Unique Responses per Power Plant """

# A participant may own/operate a Power Plant with multiple Units, and may have provided the same response for all Units at
# that facility. This version only prints unique responses.

for question_num, participant_response in zip(questions_in_base_section, dict_responses.values()):

    # Question
    print('')
    print(df_questions.iloc[question_num, 0] + ' - ' + df_questions.iloc[question_num, 1])  # Prints Question Number and Question
    print('')

    " Create DataFrame of just Power Plant, Unit, and Question "
    df_unit = df_responses_concat.iloc[:, [0, 3, question_num]].copy()
    print(tabulate(df_unit.head(2), headers='keys', tablefmt='psql', numalign='right', showindex=False))

    # Count number of unique responses across all Units
    df_unit_count = df_unit[str(df_unit.columns(2))].value_counts().rename_axis(str(df_unit.columns[2])).to_frame(
        'Count').reset_index().sort_values(by='Count', inplace=True)
    print('Unit-level Summary: ')
    print(tabulate(df_unit_count, headers='keys', tablefmt='psql', numalign='right', showindex=False))

    " Create DataFrame of Just Unique Responses per Power Plant "
    dict_unique_pp_response = {}
    for power_plant in np.arange(len(df_unit)):
        if (str(df_unit.iloc[power_plant, 0]), str(df_unit.iloc[power_plant, 2])) not in dict_unique_pp_response.items():
            dict_unique_pp_response[str(df_unit.iloc[power_plant, 0])] = str(df_unit.iloc[power_plant, 2])
        else:
            pass

    # Turn Dictionary to DataFrame
    df_pp = pd.DataFrame({str(df_unit.columns[0]): dict_unique_pp_response.keys(), str(df_unit.columns[2]):
        dict_unique_pp_response.values()})

    # Count number of unique responses per power plant
    df_pp_count = df_pp[str(df_pp.columns[1])].value_counts().rename_acis(str(df_pp.columns[1])).to_frame('Count').reset_index(
        ).sort_values(by='Count', in_place=True)
    print('Power Plant-level Summary: ')
    print(tabulate(df_pp_count, headers='keys', tablefmt='psql', numalign='right', showindex=False))

    for participant, participant_response in dict_responses.items():

        # Participant
        print(participant)

        for unique_power_plant, unique_power_plant_response in dict_unique_pp_response.items():
            # Power Plant
            print(unique_power_plant)
            # Response
            print(unique_power_plant_response)
        print('')


""" ########################################################################################################################## """
""" Set Up Functions to Format Word Document """
""" ########################################################################################################################## """


def docx_marings(doc):
    sections = doc.sections

    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def docx_add_page_number(run):
    fldchar1 = create_element('w:fldChar')
    create_attribute(fldchar1, 'w:fldCharType', 'begin')

    instrtext = create_element('w:instrText')
    create_attribute(instrtext, 'xml:space', 'preserve')
    instrtext.text = 'PAGE'

    fldchar2 = create_element('w:fldChar')
    create_attribute(fldchar2, 'w:fldCharType', 'end')

    run._r.append(fldchar1)
    run._r.append(instrtext)
    run._r.append(fldchar2)


def docx_level1_heading(dcoument_name, heading: str):
    l1 = dcoument_name.add_heading(level=1)
    run_l1 = l1.add_run(heading)  # Heading title
    run_l1.font.name = 'Times New Roman'
    run_l1.font.size = (Pt(16))
    run_l1.font.color.rgb = RGBColor.from_string('000000')


def docx_level2_heading(dcoument_name, heading: str):
    l2 = dcoument_name.add_heading(level=2)
    run_l2 = l2.add_run(heading)  # Heading title
    run_l2.font.name = 'Times New Roman'
    run_l2.font.size = (Pt(13))
    run_l2.font.color.rgb = RGBColor.from_string('000000')


def docx_level3_heading(dcoument_name, heading: str):
    l3 = dcoument_name.add_heading(level=3)
    run_l3 = l3.add_run(heading)  # Heading title
    run_l3.font.name = 'Times New Roman'
    run_l3.font.size = (Pt(13))
    run_l3.font.color.rgb = RGBColor.from_string('000000')
    l3.paragraph_format.left_indent = Inches(0.50)


def docx_level4_heading(dcoument_name, heading: str):
    l4 = dcoument_name.add_heading(level=4)
    run_l4 = l4.add_run(heading)  # Heading title
    run_l4.font.name = 'Times New Roman'
    run_l4.font.size = (Pt(13))
    run_l4.font.color.rgb = RGBColor.from_string('000000')
    l4.paragraph_format.left_indent = Inches(1.0)


def summary_by_unit(df, question_num: int):
    # Create DataFrame of just Power Plant, Unit, and Question
    df_unit = df.iloc[:, [0, 3, question_num]].copy()
    print(tabulate(df_unit, headers='keys', tablefmt='psql', numalign='right', showindex=False))

    # Counts number of unique responses per response
    df_unit_count = df_unit[str(df_unit.columns[2])].value_counts().rename_axis(str(df_unit.columns[2])).to_frame(
        'Count').reset_index().sort_values(by='Count', ascending=False, inplace=True)
    print('Unit-level Sumamry: ')
    print(tabulate(df_unit_count, headers='keys', tablefmt='psql', numalign='right', showindex=False))

    return df_unit, df_unit_count


def summary_by_power_plant(df_unit):
    # Created DataFrame of just unique Power Plant responses

    # Blank Dictionary for unique responses across Power Plants
    dict_unique_pp_response = {}
    for power_plant in np.arange(len(df_unit)):
        if (str(df_unit.iloc[power_plant, 0]), str(df_unit.iloc[power_plant, 2])) not in dict_unique_pp_response.items():
            dict_unique_pp_response[str(df_unit.iloc[power_plant, 0])] = str(df_unit.iloc[power_plant, 2])
        else:
            pass

    # Turn Dictionary to DataFrame
    df_pp = pd.DataFrame({str(df_unit.columns[0]): dict_unique_pp_response.keys(), str(df_unit.columns[2]):
        dict_unique_pp_response.values()})
    # print(tabulate(df_pp, headers='keys', tablefmt='psql', numalign='right', showindex=False))

    df_pp_count = df_pp[str(df_pp.columns[1])].value_counts().rename_axis(str(df_pp.columns[1])).to_frame('Count').reset_index()
    print('Ppwer Plant-level Summary: ')
    print(tabulate(df_pp_count, headers='keys', tablefmt='psql', numalign='right', showindex=False))

    return df_pp_count, dict_unique_pp_response


def docs_add_table(document_name, df):
    # Create blank table
    table = document_name.add_table(rows=df.shape[0]+1, cols=len(df.olumns))

    # Add headers
    for header in range(df.shape[-1]):
        table.cell(0, header).text = df.columns[header]

    # Add data/rows
    for row in range(df.shape[0]):
        for header in range(df.shape[-1]):
            table.cell(row + 1, header).text = str(df.values[row, header])


def docx_write_response_to_doc(document_name, section_name: list, level_num: int, question_set, full_response_df,
                               responses: dict, indent_start: float):

    for question_num in section_name:

        # Question
        indent_start += 0.50
        p_question = document_name.add_heading(level = level_num)
        run_question = p_question.add_run(str(question_set.iloc[question_num, 0]) + ': ' + str(question_set.iloc[question_num, 1]))
        run_question.bold = True
        run_question.font.name = 'Times New Roman'
        run_question.font.size = Pt(13)
        run_question.font.color.rgb = RGBColor.from_string('000000')
        p_question.paragraph_format.left_indet = Inches(indent_start)
        indent_start -= 0.50

        # Calculate summary statistics by Unit and Power Plant
        df_for_summary_by_power_plant_function, df_unit_summary = summary_by_unit(df_responses_concat, question_num)
        df_power_plant_summary, dict_unique = summary_by_power_plant(df_for_summary_by_power_plant_function)

        for participant, participant_response in responses.items():

            # Participant
            indent_start += 1.0
            p_entity = document_name.add_paragraph()
            p_entity.add_run(participant).bold = True
            p_entity.paragraph_format.left_indent = Inches(indent_start)
            indent_start -= 1.0

            for unique_power_plant, unique_power_plant_response in dict_unique.items():

                # Power Plant
                indent_start += 1.5
                p_power_plant = document_name.add_paragraph()
                p_power_plant.add_run(unique_power_plant).bold = True
                p_power_plant.paragraph_format.left_indent = Inches(indent_start)
                indent_start -= 1.5

                # Power Plant Response
                indent_start += 2.0
                p_response = document_name.add_paragraph()
                p_response.add_run(unique_power_plant_response)
                p_response.paragraph_format.left_indent = Inches(indent_start)
                indent_start -= 2.0

        document_name.add_paragraph()

    return document_name


def docx_style(document_name):
    # Change font style of all paragraphs and runs (overwrites settings for headers)
    for paragraph in document_name.paragraphs:
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.italic = False
        for run in paragraph.runs:
            run.font.size = Pt(13)
            run.font.name = 'Times New Roman'
            run.italic = False




